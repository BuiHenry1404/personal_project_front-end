from docx import Document

# Tạo tài liệu Word mới
doc = Document()
doc.add_heading("Cấu trúc Cơ sở Dữ liệu: Hệ thống quản lý nhà trọ", level=1)

# Danh sách bảng và cột
tables = [
    {
        "name": "user",
        "description": "Thông tin của user",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "fullname", "type": "VARCHAR(100)", "note": "Họ tên người dùng"},
            {"name": "email", "type": "VARCHAR(100)", "note": "Email người dùng"},
            {"name": "cccd", "type": "VARCHAR(20)", "note": "Căn cước công dân"},
            {"name": "phone_number", "type": "VARCHAR(20)", "note": "Số điện thoại"},
            {"name": "permanent_address", "type": "TEXT", "note": "Địa chỉ thường trú"},
            {"name": "date_of_birth", "type": "DATE", "note": "Ngày sinh"},
            {"name": "recorded_at", "type": "DATE", "note": "Ngày nhập trọ"},
            {"name": "license_plate_number", "type": "VARCHAR(20)", "note": "Biển số xe"},
            {"name": "note", "type": "TEXT", "note": "Ghi chú"},
            {"name": "username", "type": "VARCHAR(50)", "note": "Tên đăng nhập, duy nhất"},
            {"name": "password", "type": "VARCHAR(255)", "note": "Mật khẩu đã mã hóa"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật"}
        ]
    },
    {
        "name": "role",
        "description": "Vai trò người dùng",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "name", "type": "VARCHAR(50)", "note": "Tên vai trò"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật"}
        ]
    },
    {
        "name": "checkout_request",
        "description": "Yêu cầu trả phòng",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "room_id", "type": "INT", "note": "Mã phòng, không được null"},
            {"name": "user_id", "type": "INT", "note": "Mã người dùng, không được null"},
            {"name": "request_date", "type": "DATE", "note": "Ngày yêu cầu, không được null"},
            {"name": "status", "type": "VARCHAR(50)", "note": "Trạng thái yêu cầu, không được null"},
            {"name": "reason", "type": "TEXT", "note": "Lý do trả phòng"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "contract",
        "description": "Hợp đồng thuê trọ",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "room_id", "type": "INT", "note": "Mã phòng, không được null"},
            {"name": "user_id", "type": "INT", "note": "Mã người dùng, không được null"},
            {"name": "start_date", "type": "DATE", "note": "Ngày bắt đầu, không được null"},
            {"name": "end_date", "type": "DATE", "note": "Ngày kết thúc"},
            {"name": "deposit", "type": "DECIMAL(10, 2)", "note": "Tiền đặt cọc, không được null"},
            {"name": "monthly_rent", "type": "DECIMAL(10, 2)", "note": "Tiền thuê hàng tháng, không được null"},
            {"name": "status", "type": "VARCHAR(50)", "note": "Trạng thái hợp đồng, không được null"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "maintenance_fee",
        "description": "Dịch vụ bảo trì",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "maintenance_request_id", "type": "INT", "note": "Mã yêu cầu bảo trì, không được null"},
            {"name": "price", "type": "DECIMAL(10, 2)", "note": "Giá dịch vụ, không được null"},
            {"name": "total_fee", "type": "DECIMAL(10, 2)", "note": "Tổng phí"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "maintenance_request",
        "description": "Yêu cầu bảo trì",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "contract_id", "type": "INT", "note": "Mã hợp đồng, không được null"},
            {"name": "service_room_id", "type": "INT", "note": "Mã dịch vụ phòng, không được null"},
            {"name": "status", "type": "VARCHAR(50)", "note": "Trạng thái yêu cầu, không được null"},
            {"name": "request_date", "type": "DATE", "note": "Ngày yêu cầu, không được null"},
            {"name": "decision", "type": "TEXT", "note": "Quyết định xử lý"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "service",
        "description": "Dịch vụ trong quá trình thuê trọ",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "name", "type": "VARCHAR(100)", "note": "Tên dịch vụ, không được null"},
            {"name": "price", "type": "DECIMAL(10, 2)", "note": "Giá dịch vụ, không được null"},
            {"name": "is_active", "type": "TINYINT(1)", "note": "Trạng thái hoạt động, mặc định là 1"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "service_room",
        "description": "Bảng trung gian giữa phòng và dịch vụ miêu tả dịch vụ có trong phòng",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "room_id", "type": "INT", "note": "Mã phòng, không được null"},
            {"name": "service_id", "type": "INT", "note": "Mã dịch vụ, không được null"},
            {"name": "quantity", "type": "INT", "note": "Số lượng, không được null"},
            {"name": "price", "type": "DECIMAL(10, 2)", "note": "Giá, không được null"},
            {"name": "is_active", "type": "TINYINT(1)", "note": "Trạng thái hoạt động, mặc định là 1"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "room",
        "description": "Phòng trọ",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "number", "type": "VARCHAR(50)", "note": "Số phòng, không được null"},
            {"name": "price", "type": "DECIMAL(10, 2)", "note": "Giá phòng, không được null"},
            {"name": "length", "type": "DECIMAL(5, 2)", "note": "Chiều dài phòng"},
            {"name": "width", "type": "DECIMAL(5, 2)", "note": "Chiều rộng phòng"},
            {"name": "status", "type": "VARCHAR(50)", "note": "Trạng thái phòng, không được null"},
            {"name": "note", "type": "TEXT", "note": "Ghi chú"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "user_role",
        "description": "Bảng trung gian giữa user và role thể hiện vai trò của người dùng",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "user_id", "type": "INT", "note": "Mã người dùng, không được null"},
            {"name": "role_id", "type": "INT", "note": "Mã vai trò, không được null"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "utility_index",
        "description": "Miêu tả chỉ số điện nước hàng tháng của phòng",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "room_id", "type": "INT", "note": "Mã phòng, không được null"},
            {"name": "electric_old_index", "type": "INT", "note": "Chỉ số điện cũ, không được null"},
            {"name": "electric_new_index", "type": "INT", "note": "Chỉ số điện mới, không được null"},
            {"name": "electric_usage", "type": "INT", "note": "Lượng điện sử dụng, không được null"},
            {"name": "electric_recorded_at", "type": "DATE", "note": "Ngày ghi chỉ số điện, không được null"},
            {"name": "water_old_index", "type": "INT", "note": "Chỉ số nước cũ, không được null"},
            {"name": "water_new_index", "type": "INT", "note": "Chỉ số nước mới, không được null"},
            {"name": "water_usage", "type": "INT", "note": "Lượng nước sử dụng, không được null"},
            {"name": "water_recorded_at", "type": "DATE", "note": "Ngày ghi chỉ số nước, không được null"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "invoice_detail",
        "description": "Chi tiết hóa đơn tiền nhà",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "invoice_id", "type": "INT", "note": "Mã hóa đơn, không được null"},
            {"name": "description", "type": "TEXT", "note": "Mô tả"},
            {"name": "quantity", "type": "INT", "note": "Số lượng, không được null"},
            {"name": "price", "type": "DECIMAL(10, 2)", "note": "Giá, không được null"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "invoice",
        "description": "Hóa đơn tiền nhà",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "phone_number", "type": "VARCHAR(20)", "note": "Số điện thoại, không được null"},
            {"name": "cccd", "type": "VARCHAR(20)", "note": "Căn cước công dân, không được null"},
            {"name": "email", "type": "VARCHAR(100)", "note": "Email, không được null"},
            {"name": "user_id", "type": "INT", "note": "Mã người dùng, không được null"},
            {"name": "room_id", "type": "INT", "note": "Mã phòng, không được null"},
            {"name": "note", "type": "TEXT", "note": "Ghi chú"},
            {"name": "total_price", "type": "DECIMAL(10, 2)", "note": "Tổng tiền, không được null"},
            {"name": "due_date", "type": "DATE", "note": "Ngày đến hạn, không được null"},
            {"name": "payment_date", "type": "DATE", "note": "Ngày thanh toán"},
            {"name": "status", "type": "VARCHAR(50)", "note": "Trạng thái hóa đơn, không được null"},
            {"name": "payment_id", "type": "INT", "note": "Mã thanh toán"},
            {"name": "invoice_detail_id", "type": "INT", "note": "Mã chi tiết hóa đơn"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "room_member",
        "description": "Thông tin người trong phòng là bảng trung gian giữa user và room",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "room_id", "type": "INT", "note": "Mã phòng, không được null"},
            {"name": "user_id", "type": "INT", "note": "Mã người dùng, không được null"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    },
    {
        "name": "payment",
        "description": "Phương thức thanh toán",
        "columns": [
            {"name": "id", "type": "INT", "note": "Khóa chính, tự tăng"},
            {"name": "method", "type": "VARCHAR(50)", "note": "Phương thức thanh toán, không được null"},
            {"name": "status", "type": "VARCHAR(50)", "note": "Trạng thái thanh toán, không được null"},
            {"name": "created_at", "type": "TIMESTAMP", "note": "Ngày tạo, mặc định thời gian hiện tại"},
            {"name": "updated_at", "type": "TIMESTAMP", "note": "Ngày cập nhật, tự động cập nhật"}
        ]
    }
]

# Thêm nội dung bảng vào Word
for table in tables:
    doc.add_heading(f"Bảng: {table['name']}", level=2)
    doc.add_paragraph(table["description"])
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'  # Thêm style cho bảng
    
    # Thiết lập header
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'Tên cột'
    hdr_cells[1].text = 'Kiểu dữ liệu'
    hdr_cells[2].text = 'Ghi chú'
    
    # Thêm nội dung từng cột
    for col in table["columns"]:
        row_cells = t.add_row().cells
        row_cells[0].text = col["name"]
        row_cells[1].text = col["type"]
        row_cells[2].text = col["note"]
    
    # Thêm khoảng trống sau mỗi bảng
    doc.add_paragraph("")

# Lưu tài liệu
output_path = "D:/persornal_project/csdl_thue_tro.docx"
doc.save(output_path)

print(f"File đã được tạo tại: {output_path}")
